"""
日报预测 API 服务

启动: python api_server.py
接口: POST /api/daily/upload-and-predict
"""
import os
import sys
import re
import json
import tempfile
import shutil
from datetime import timedelta
from io import BytesIO

import numpy as np
import pandas as pd
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn

# 确保 src 在 path 中
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SRC_DIR = os.path.join(BASE_DIR, "src")
sys.path.insert(0, SRC_DIR)

from process_daily import parse_one_report, aggregate_to_weekly
from predict_api import predict_from_dataframe
from contract_extract import extract_contract_fields, format_result

from datetime import datetime
from fastapi import Depends, HTTPException, status
from auth import (
    hash_password, verify_password,
    create_access_token, create_refresh_token, decode_token,
)
from auth_middleware import get_current_user, require_admin
from database import get_db
from models import (
    UserLogin, UserCreate, UserUpdate, UserOut,
    TokenResponse, RefreshRequest,
    ChangePasswordRequest, ResetPasswordRequest,
)

app = FastAPI(title="日报预测API", version="1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ========== 认证路由 ==========

@app.post("/api/auth/login", response_model=TokenResponse)
def login(body: UserLogin):
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT * FROM users WHERE username = ? AND is_active = 1",
            (body.username,)
        ).fetchone()
        if row is None or not verify_password(body.password, row["password"]):
            raise HTTPException(status_code=401, detail="用户名或密码错误")

        user = dict(row)
        access_token = create_access_token(user["id"], user["role"])
        refresh_token = create_refresh_token(user["id"], user["token_version"])

        conn.execute(
            "UPDATE users SET last_login = datetime('now', 'localtime') WHERE id = ?",
            (user["id"],)
        )
        conn.commit()

        return {
            "access_token": access_token,
            "refresh_token": refresh_token,
            "user": UserOut(**user).model_dump(),
        }
    finally:
        conn.close()


@app.post("/api/auth/refresh")
def refresh_token(body: RefreshRequest):
    payload = decode_token(body.refresh_token, "refresh")
    if payload is None:
        raise HTTPException(status_code=401, detail="登录已过期，请重新登录")

    user_id = int(payload["sub"])
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT * FROM users WHERE id = ? AND is_active = 1",
            (user_id,)
        ).fetchone()
        if row is None:
            raise HTTPException(status_code=401, detail="账号已被禁用或删除")
        if row["token_version"] != payload.get("ver"):
            raise HTTPException(status_code=401, detail="密码已重置，请重新登录")

        new_access = create_access_token(user_id, row["role"])
        return {"access_token": new_access}
    finally:
        conn.close()


@app.get("/api/auth/me", response_model=UserOut)
def me(user: dict = Depends(get_current_user)):
    return UserOut(**user).model_dump()


@app.post("/api/auth/change-password")
def change_password(body: ChangePasswordRequest, user: dict = Depends(get_current_user)):
    if not verify_password(body.old_password, user["password"]):
        raise HTTPException(status_code=400, detail="原密码错误")

    new_hash = hash_password(body.new_password)
    new_version = user["token_version"] + 1

    conn = get_db()
    try:
        conn.execute(
            "UPDATE users SET password = ?, token_version = ? WHERE id = ?",
            (new_hash, new_version, user["id"]),
        )
        conn.commit()
        return {"message": "密码修改成功，请重新登录"}
    finally:
        conn.close()


def _dify_extract_pdf(file_bytes: bytes, filename: str) -> str:
    """通过 Dify AI 服务提取扫描件 PDF 中的文字。"""
    import json as _json
    import urllib.request as _req
    import urllib.error as _err

    DIFY_BASE = "http://172.16.204.124/v1"
    DIFY_TOKEN = "Bearer app-eqYhrSTYq86vbBXnAAuMQM15"

    # Step 1: 上传文件到 Dify
    boundary = "----DifyUploadBoundary2025"
    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'
        f"Content-Type: application/pdf\r\n\r\n"
    ).encode("utf-8") + file_bytes + (
        f"\r\n--{boundary}\r\n"
        f'Content-Disposition: form-data; name="user"\r\n\r\n'
        f"admin\r\n"
        f"--{boundary}--\r\n"
    ).encode("utf-8")

    upload_req = _req.Request(
        f"{DIFY_BASE}/files/upload",
        data=body,
        headers={
            "Authorization": DIFY_TOKEN,
            "Content-Type": f"multipart/form-data; boundary={boundary}",
        },
        method="POST",
    )
    try:
        with _req.urlopen(upload_req, timeout=30) as resp:
            upload_result = _json.loads(resp.read().decode("utf-8"))
        file_id = upload_result.get("id") or upload_result.get("data", {}).get("id", "")
    except _err.URLError as e:
        raise RuntimeError(f"Dify 文件上传失败: {e}")

    if not file_id:
        raise RuntimeError("Dify 文件上传未返回 file_id")

    # Step 2: 调用 Chat API 提取文本
    chat_body = _json.dumps({
        "inputs": {
            "contract_file": {
                "transfer_method": "local_file",
                "upload_file_id": file_id,
                "type": "document",
            }
        },
        "query": "请提取这份合同文件的全部文字内容，包括合同编号、工程名称、发包人、承包人、开工日期、竣工日期、合同价款、工程地点、工程范围等所有条款。",
        "response_mode": "blocking",
        "user": "admin",
    }).encode("utf-8")

    chat_req = _req.Request(
        f"{DIFY_BASE}/chat-messages",
        data=chat_body,
        headers={
            "Authorization": DIFY_TOKEN,
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with _req.urlopen(chat_req, timeout=120) as resp:
            chat_result = _json.loads(resp.read().decode("utf-8"))
    except _err.URLError as e:
        raise RuntimeError(f"Dify 文档识别失败: {e}")

    answer = chat_result.get("answer") or chat_result.get("data", {}).get("answer", "")
    if not answer.strip():
        raise RuntimeError("Dify 未返回文字内容")

    return answer


def _extract_docx_images(file_bytes: bytes) -> str:
    """从 DOCX 中提取嵌入图片，通过 Dify 批量 OCR 识别文字。"""
    import zipfile
    import json as _json
    import urllib.request as _req
    import urllib.error as _err

    DIFY_BASE = "http://172.16.204.124/v1"
    DIFY_TOKEN = "Bearer app-eqYhrSTYq86vbBXnAAuMQM15"

    # 解压 DOCX 获取图片
    image_data = []
    with zipfile.ZipFile(BytesIO(file_bytes)) as zf:
        for name in sorted(zf.namelist()):
            if name.startswith("word/media/") and name.lower().endswith(
                (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")
            ):
                image_data.append((name, zf.read(name)))

    if not image_data:
        return ""  # 没有嵌入图片

    # 上传所有图片到 Dify，获取 file_id 列表
    file_ids = []
    for img_name, img_bytes in image_data:
        short_name = img_name.split("/")[-1]
        boundary = "----DifyImgUpload2025"
        body = (
            f"--{boundary}\r\n"
            f'Content-Disposition: form-data; name="file"; filename="{short_name}"\r\n'
            f"Content-Type: image/png\r\n\r\n"
        ).encode("utf-8") + img_bytes + (
            f"\r\n--{boundary}\r\n"
            f'Content-Disposition: form-data; name="user"\r\n\r\n'
            f"admin\r\n"
            f"--{boundary}--\r\n"
        ).encode("utf-8")

        upload_req = _req.Request(
            f"{DIFY_BASE}/files/upload",
            data=body,
            headers={
                "Authorization": DIFY_TOKEN,
                "Content-Type": f"multipart/form-data; boundary={boundary}",
            },
            method="POST",
        )
        try:
            with _req.urlopen(upload_req, timeout=30) as resp:
                result = _json.loads(resp.read().decode("utf-8"))
            fid = result.get("id") or result.get("data", {}).get("id", "")
            if fid:
                file_ids.append(fid)
        except _err.URLError:
            continue

    if not file_ids:
        raise RuntimeError("DOCX 图片上传 Dify 失败")

    # 构建带多张图片的 Chat 请求
    inputs = {}
    for i, fid in enumerate(file_ids[:10]):  # 最多 10 张图
        inputs[f"image_{i+1}"] = {
            "transfer_method": "local_file",
            "upload_file_id": fid,
            "type": "image",
        }

    chat_body = _json.dumps({
        "inputs": inputs,
        "query": "请逐一识别这些图片中的全部文字内容，包括合同编号、工程名称、发包人、承包人、开工日期、竣工日期、合同价款、工程地点、工程范围等所有条款。",
        "response_mode": "blocking",
        "user": "admin",
    }).encode("utf-8")

    chat_req = _req.Request(
        f"{DIFY_BASE}/chat-messages",
        data=chat_body,
        headers={
            "Authorization": DIFY_TOKEN,
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with _req.urlopen(chat_req, timeout=120) as resp:
            chat_result = _json.loads(resp.read().decode("utf-8"))
    except _err.URLError as e:
        raise RuntimeError(f"Dify 图片 OCR 失败: {e}")

    answer = chat_result.get("answer") or chat_result.get("data", {}).get("answer", "")
    return answer


# ============================================================
# 1. 文本提取
# ============================================================

def extract_text(file_bytes: bytes, filename: str) -> str:
    """从上传文件中提取纯文本"""
    ext = os.path.splitext(filename)[1].lower()

    if ext in (".txt", ".text"):
        return file_bytes.decode("utf-8", errors="ignore")

    if ext in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(BytesIO(file_bytes))
            # 同时提取段落和表格中的文本
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
            text = "\n".join(parts)

            # 如果文本为空，尝试提取文档中的图片并 OCR
            if not text.strip():
                img_text = _extract_docx_images(file_bytes)
                if img_text.strip():
                    text = "[以下内容由 Dify AI 从文档图片中识别]\n\n" + img_text
                else:
                    raise RuntimeError(
                        "Word 文档中既无文字也无图片，或图片 OCR 识别失败。"
                        "请确认文档内容可见（非空白页），或将合同另存为 PDF 后重新上传。"
                    )
            return text
        except ImportError:
            raise RuntimeError("请安装 python-docx")
        except Exception as e:
            raise RuntimeError(f"Word 文档解析失败: {str(e)}")

    if ext == ".pdf":
        text = ""
        # 方案1: pdfplumber（优先，对中文支持好）
        try:
            import pdfplumber
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
        except ImportError:
            raise RuntimeError("请安装 pdfplumber")
        except Exception:
            text = ""

        # 方案2: pdfminer.six 兜底
        if not text.strip():
            try:
                from pdfminer.high_level import extract_text as pdfminer_extract
                text = pdfminer_extract(BytesIO(file_bytes))
            except Exception:
                pass

        # 方案3: pymupdf (fitz) 兜底
        if not text.strip():
            try:
                import fitz
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page in doc:
                    t = page.get_text()
                    if t:
                        text += t + "\n"
                doc.close()
            except Exception:
                pass

        # 方案4: Dify AI 文档识别（处理扫描件，速度快）
        if not text.strip():
            try:
                text = _dify_extract_pdf(file_bytes, filename)
            except Exception as e:
                _ocr_error = str(e)

        if not text.strip():
            _ocr_detail = f" Dify 错误: {_ocr_error}" if _ocr_error else ""
            raise RuntimeError(
                "无法从 PDF 中提取文本（已尝试 pdfplumber / pdfminer / pymupdf / Dify AI 四种方式均失败）。" +
                _ocr_detail +
                " 建议将 PDF 转换为 Word 或 TXT 格式后重新上传。"
            )
        return text

    if ext in (".xlsx", ".xls"):
        try:
            dfs = pd.read_excel(BytesIO(file_bytes), sheet_name=None)
            text = ""
            for sheet, df in dfs.items():
                text += df.to_string(index=False) + "\n"
            return text
        except Exception as e:
            raise RuntimeError(f"Excel 解析失败: {str(e)}")

    raise RuntimeError(f"不支持的文件格式: {ext}")


# ============================================================
# 2. 按日拆分
# ============================================================

def split_by_date(text: str) -> list:
    """
    按日期标记拆分为逐日报告。
    日期间隔：跨天数 > 0 则视为新一天。
    支持格式: 2025-04-12, 2025-4-7, 2025年4月12日
    """
    date_pat = re.compile(
        r"(\d{4})[-/年](\d{1,2})[-/月](\d{1,2})[日]?"
    )
    matches = list(date_pat.finditer(text))

    if len(matches) < 2:
        return [text.strip()] if text.strip() else []

    segments = []
    for i, m in enumerate(matches):
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        day_text = text[start:end].strip()
        if day_text:
            segments.append(day_text)

    return segments


# ============================================================
# 3. 核心流程
# ============================================================

def process_uploaded_file(file_bytes: bytes, filename: str, project_code: str):
    """完整处理链路：文本→解析→聚合→预测用DataFrame，同时返回最后一天原始文本"""
    # 提取文本
    text = extract_text(file_bytes, filename)

    # 按日拆分
    daily_texts = split_by_date(text)
    if not daily_texts:
        raise RuntimeError("未能从文件中识别出日报内容")

    # 保留最后一天的原始文本
    last_day_text = daily_texts[-1].strip() if daily_texts else ""

    # 逐日解析
    records = []
    for day_text in daily_texts:
        rec = parse_one_report(day_text, filename)
        if rec["date"] is not None:
            records.append(rec)

    if not records:
        raise RuntimeError("未能从日报中提取到日期信息")

    daily_df = pd.DataFrame(records)
    daily_df = daily_df.dropna(subset=["date"]).sort_values("date")

    # 聚合为周级
    weekly_df = aggregate_to_weekly(daily_df, project_code)
    if len(weekly_df) < 4:
        raise RuntimeError(f"仅提取到 {len(weekly_df)} 周数据，至少需要4周")

    return weekly_df, last_day_text


# ============================================================
# 4. JSON 序列化工具
# ============================================================

def _sanitize_value(v):
    """递归转换 numpy 类型为原生 Python 类型，确保 JSON 可序列化。"""
    if v is None:
        return v
    if isinstance(v, (np.integer,)):
        return int(v)
    if isinstance(v, (np.floating,)):
        if np.isnan(v) or np.isinf(v):
            return 0
        return float(v)
    if isinstance(v, np.ndarray):
        return v.tolist()
    if isinstance(v, (datetime, pd.Timestamp)):
        return str(v)
    if hasattr(v, "isoformat"):
        return str(v)
    if isinstance(v, dict):
        return {k: _sanitize_value(val) for k, val in v.items()}
    if isinstance(v, (list, tuple)):
        return [_sanitize_value(item) for item in v]
    return v


# ============================================================
# 5. API 端点
# ============================================================

@app.post("/api/daily/upload-and-predict")
async def upload_and_predict(
    file: UploadFile = File(...),
    predict_weeks: int = Form(1),
    current_user: dict = Depends(get_current_user),
):
    """上传日报文件，解析并预测"""
    try:
        # 读取文件
        file_bytes = await file.read()
        project_code = "UPLOAD"

        # 处理
        weekly_df, last_day_text = process_uploaded_file(file_bytes, file.filename, project_code)

        # 预测（传入最后一天原文用于生成日报文本）
        prediction = predict_from_dataframe(weekly_df, predict_weeks, last_day_text)

        # 组装解析结果
        weeks_data = weekly_df.to_dict(orient="records")
        # 转换日期类型
        for w in weeks_data:
            for k, v in w.items():
                if hasattr(v, "isoformat"):
                    w[k] = str(v)
                elif isinstance(v, (np.integer,)):
                    w[k] = int(v)
                elif isinstance(v, (np.floating,)):
                    w[k] = float(v)
                elif pd.isna(v):
                    w[k] = 0

        return {
            "parsed": {
                "weeks": weeks_data,
                "week_count": len(weekly_df)
            },
            "prediction": _sanitize_value(prediction)
        }

    except RuntimeError as e:
        return {
            "parsed": None,
            "prediction": None,
            "error": str(e)
        }
    except Exception as e:
        return {
            "parsed": None,
            "prediction": None,
            "error": f"处理异常: {str(e)}"
        }


class PredictFromDataRequest(BaseModel):
    weeks: list
    predict_weeks: int = 1


@app.post("/api/daily/predict-from-data")
async def predict_from_data(
    body: PredictFromDataRequest,
    current_user: dict = Depends(get_current_user),
):
    """基于已上传并解析的周级数据直接预测（无需重新上传文件）。"""
    try:
        weeks = body.weeks
        predict_weeks = body.predict_weeks
        if not weeks or len(weeks) < 4:
            raise RuntimeError(f"数据不足（{len(weeks) if weeks else 0}周），至少需要4周")

        weekly_df = pd.DataFrame(weeks)
        if "week_start" in weekly_df.columns:
            weekly_df["week_start"] = pd.to_datetime(weekly_df["week_start"])

        prediction = predict_from_dataframe(weekly_df, predict_weeks, last_day_text="")
        return _sanitize_value(prediction)

    except RuntimeError as e:
        return {"error": str(e)}
    except Exception as e:
        return {"error": f"处理异常: {str(e)}"}


# ============================================================
# 合同信息提取
# ============================================================

@app.post("/api/contract/extract")
async def contract_extract(
    file: UploadFile = File(...),
    use_ai: bool = Form(False),
    current_user: dict = Depends(get_current_user),
):
    """上传合同文件，提取关键信息（项目名称、日期、金额、工期等）。"""
    try:
        file_bytes = await file.read()
        text = extract_text(file_bytes, file.filename)

        if not text or not text.strip():
            raise RuntimeError("未能从文件中提取出文本内容（文件可能是纯图片格式，无文字层）")

        raw_result = extract_contract_fields(text)
        result = format_result(raw_result)

        # 统计命中字段数 + 提取关键摘要
        fields_found = sum(1 for v in result.values() if v.get("value") is not None)
        project_name = result.get("project_name", {}).get("value") or ""
        contract_no = result.get("contract_no", {}).get("value") or ""
        start_date = result.get("start_date", {}).get("value") or ""
        end_date = result.get("end_date", {}).get("value") or ""

        # 保存到历史记录
        conn = get_db()
        try:
            conn.execute(
                "INSERT INTO contract_records (filename, file_size, project_name, contract_no, "
                "start_date, end_date, text_length, fields_json, status, created_at) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'success', datetime('now', 'localtime'))",
                (file.filename, len(file_bytes), project_name, contract_no,
                 start_date, end_date, len(text),
                 json.dumps(result, ensure_ascii=False))
            )
            conn.commit()
        finally:
            conn.close()

        return {
            "file_name": file.filename,
            "file_size": len(file_bytes),
            "text_length": len(text),
            "text_preview": text[:500] if text else "",
            "fields_found": fields_found,
            "fields": result,
        }

    except RuntimeError as e:
        return {
            "file_name": file.filename if file else None,
            "fields": None,
            "error": str(e),
        }
    except Exception as e:
        return {
            "file_name": file.filename if file else None,
            "fields": None,
            "error": f"处理异常: {str(e)}",
        }


@app.get("/api/contract/history")
def contract_history(
    keyword: str = "",
    current_user: dict = Depends(get_current_user),
):
    """返回历史上传合同列表，支持按文件名/项目名称/合同编号搜索。"""
    conn = get_db()
    try:
        if keyword:
            kw = f"%{keyword}%"
            rows = conn.execute(
                "SELECT id, filename, file_size, project_name, contract_no, start_date, end_date, "
                "text_length, status, error_msg, created_at "
                "FROM contract_records WHERE filename LIKE ? OR project_name LIKE ? OR contract_no LIKE ? "
                "ORDER BY id DESC LIMIT 50",
                (kw, kw, kw)
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT id, filename, file_size, project_name, contract_no, start_date, end_date, "
                "text_length, status, error_msg, created_at "
                "FROM contract_records ORDER BY id DESC LIMIT 50"
            ).fetchall()
        return {"items": [dict(r) for r in rows]}
    finally:
        conn.close()


@app.get("/api/contract/history/{record_id}")
def contract_history_detail(record_id: int, current_user: dict = Depends(get_current_user)):
    """返回某次历史记录对应的完整提取结果。"""
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT * FROM contract_records WHERE id = ?", (record_id,)
        ).fetchone()
        if row is None:
            return {"error": "记录不存在"}
        data = dict(row)
        data["fields"] = json.loads(data["fields_json"]) if data.get("fields_json") else None
        return data
    finally:
        conn.close()


@app.delete("/api/contract/history/{record_id}")
def contract_history_delete(record_id: int, current_user: dict = Depends(get_current_user)):
    """删除某条历史记录。"""
    conn = get_db()
    try:
        conn.execute("DELETE FROM contract_records WHERE id = ?", (record_id,))
        conn.commit()
        return {"message": "已删除"}
    finally:
        conn.close()


@app.get("/api/daily/health")
def health():
    return {"status": "ok"}


# ============================================================
# 5. 管理员路由
# ============================================================

from admin_routes import router as admin_router
app.include_router(admin_router)

# ============================================================
# 6. 启动
# ============================================================

if __name__ == "__main__":
    print("=" * 50)
    print("  日报预测 API 服务")
    print(f"  http://127.0.0.1:8000")
    print(f"  http://127.0.0.1:8000/docs")
    print("=" * 50)
    uvicorn.run(app, host="0.0.0.0", port=8000)
