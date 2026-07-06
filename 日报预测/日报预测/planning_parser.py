"""
策划书文档解析模块

解析中材国际标准总包项目管理策划书（.docx），提取：
- 项目元信息（名称、规模、工期、地点等）
- 里程碑节点（名称、日期、完成标准、分类）
- 关键路径（4条路径，含各阶段日期）
- 月度任务分解（设计/采购/施工）
"""

import re
import os
import subprocess
import tempfile
import shutil
from io import BytesIO
from datetime import datetime
from typing import Optional

from docx import Document


# ============================================================
# 常量：类别关键词映射
# ============================================================

# 里程碑分类关键词
CATEGORY_KEYWORDS = {
    "设计": ["设计", "图", "图纸", "方案", "规划"],
    "采购": ["采购", "订货", "订", "提资", "招标", "合同", "签", "中标"],
    "施工": ["施工", "开工", "安装", "土建", "封顶", "浇筑", "开挖", "钢结构", "砌筑"],
    "调试": ["调试", "试运转", "试车", "点火", "投料", "达产", "达标", "竣工", "验收", "考核"],
}

# 部门阶段关键词
PHASE_DESIGN = "设计"
PHASE_PROCUREMENT = "采购"
PHASE_CONSTRUCTION = "施工"


# ============================================================
# 日期解析
# ============================================================

def parse_date(text: str) -> Optional[str]:
    """从文本中解析日期，返回 ISO 格式 YYYY-MM-DD 或 None"""
    if not text:
        return None

    text = text.strip()

    # 支持格式:
    # 2020.12.16 / 2020-12-16 / 2020/12/16
    # 2020.12.16前 / 12月16日 / 2020年12月16日
    # 2021年春节 / 2021年2月 / 2021年8月底

    # 完整日期 YYYY.M.D 或 YYYY-M-D 或 YYYY/M/D
    m = re.search(r"(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})", text)
    if m:
        try:
            y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
            return f"{y:04d}-{mo:02d}-{d:02d}"
        except ValueError:
            pass

    # 中文日期 YYYY年M月D日
    m = re.search(r"(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日", text)
    if m:
        try:
            y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
            return f"{y:04d}-{mo:02d}-{d:02d}"
        except ValueError:
            pass

    # 年月 YYYY年M月（取当月第一天）
    m = re.search(r"(\d{4})\s*年\s*(\d{1,2})\s*月", text)
    if m:
        try:
            y, mo = int(m.group(1)), int(m.group(2))
            return f"{y:04d}-{mo:02d}-01"
        except ValueError:
            pass

    # "月中旬/月底" → 取15日/28日
    m = re.search(r"(\d{4})[.\-/](\d{1,2})\s*(月|[月初底])", text)
    if m:
        try:
            y, mo = int(m.group(1)), int(m.group(2))
            if "底" in text:
                return f"{y:04d}-{mo:02d}-28"
            return f"{y:04d}-{mo:02d}-15"
        except ValueError:
            pass

    return None


def classify_milestone(name: str) -> str:
    """根据里程碑名称关键词分类"""
    for category, keywords in CATEGORY_KEYWORDS.items():
        for kw in keywords:
            if kw in name:
                return category
    return "项目节点"


def classify_path_step(text: str) -> str:
    """将关键路径中的步骤归类到设计/采购/施工"""
    if any(kw in text for kw in ["图", "图纸", "方案", "设计", "规划"]):
        return PHASE_DESIGN
    if any(kw in text for kw in ["采购", "订货", "订", "提资", "招标", "合同", "到厂", "发货"]):
        return PHASE_PROCUREMENT
    if any(kw in text for kw in ["施工", "开工", "安装", "土建", "封顶", "浇筑", "开挖",
                                  "钢结构", "砌筑", "交付", "竣工", "调试", "试车"]):
        return PHASE_CONSTRUCTION
    return "其他"


# ============================================================
# 表格提取
# ============================================================

def extract_tables(doc: Document) -> list:
    """提取文档中所有表格的行数据"""
    tables_data = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables_data.append(rows)
    return tables_data


# ============================================================
# 项目信息提取
# ============================================================

def parse_project_info(paragraphs: list) -> dict:
    """从文档开头段落提取项目元信息"""
    info = {
        "name": "",
        "scale": "",
        "owner": "",
        "location": "",
        "contract_period": "",
        "start_date": "",
        "end_date": "",
        "planning_version": "",
    }

    full_text = "\n".join(p[:300] for p in paragraphs[:80])

    # 项目名称 — 取连续的项目描述行
    name_parts = []
    for p in paragraphs[:15]:
        t = p.strip()
        if not t or len(t) > 100:
            continue
        # 排除表头、版本号等
        if any(kw in t for kw in ["版本号", "编制", "审核", "批准", "Rev", "签字", "修改"]):
            continue
        if re.match(r"^[A-Z]+\d", t):  # 编号行
            continue
        name_parts.append(t)
        if "策划书" in t or "有限公司" in t:
            break

    info["name"] = "".join(name_parts) if name_parts else ""

    # 规模 — 匹配如 8000t/d, 5000t/d
    m = re.search(r"(\d{3,5})\s*t/\s*d", full_text)
    if m:
        info["scale"] = m.group(0).replace(" ", "")

    # 业主
    m = re.search(r"项目业主[为是：:]\s*(.+?)(?:，|。|\n|$)", full_text)
    if m:
        info["owner"] = m.group(1).strip()

    # 地点
    m = re.search(r"工程地点[：:]\s*(.+?)(?:，|。|\n|$)", full_text)
    if m:
        info["location"] = m.group(1).strip()

    # 合同工期
    m = re.search(r"合同工期[：:]\s*(.+?)(?:，|。|\n|$)", full_text)
    if m:
        info["contract_period"] = m.group(1).strip()

    # 开工/竣工日期
    m = re.search(r"开工[日期]*[：:]*\s*(\d{4}[.\-/]\d{1,2}[.\-/]\d{1,2})", full_text)
    if m:
        info["start_date"] = parse_date(m.group(0)) or ""
    m = re.search(r"(?:计划)?竣工[日期]*[：:]*\s*(\d{4}[.\-/]\d{1,2}[.\-/]\d{1,2})", full_text)
    if m:
        info["end_date"] = parse_date(m.group(0)) or ""

    # 版本号
    m = re.search(r"(?:版本号|策划版本)[：:]*\s*(.+?)(?:\n|$)", full_text)
    if m:
        info["planning_version"] = m.group(1).strip()

    return info


# ============================================================
# 里程碑解析
# ============================================================

def parse_milestones(tables_data: list) -> list:
    """从表格中识别并解析里程碑数据"""
    milestones = []

    for rows in tables_data:
        if not rows or len(rows) < 3:
            continue

        header = rows[0]
        header_text = " ".join(header)

        # 识别里程碑表格：表头含"节点名称"+"完成时间"
        has_node = any("节点" in c for c in header)
        has_date = any("时间" in c or "日期" in c for c in header)

        if not (has_node and has_date):
            continue

        # 确定列索引
        name_col = date_col = criteria_col = None
        for i, col_name in enumerate(header):
            if "节点" in col_name and "名称" in col_name:
                name_col = i
            elif "时间" in col_name or "日期" in col_name:
                date_col = i
            elif "标准" in col_name or "条件" in col_name or "要求" in col_name:
                criteria_col = i

        if name_col is None or date_col is None:
            # 回退：按位置猜测（第1列名称，第2列日期，第3列标准）
            if len(header) >= 2:
                name_col = 0
                date_col = 1
                criteria_col = 2 if len(header) > 2 else None
            else:
                continue

        # 解析数据行
        for row in rows[1:]:
            if len(row) <= name_col:
                continue
            name = row[name_col].strip()
            if not name:
                continue

            # 跳过表头重复行
            if "节点名称" in name or "里程碑" in name:
                continue

            date_str = row[date_col].strip() if len(row) > date_col else ""
            criteria = row[criteria_col].strip() if criteria_col is not None and len(row) > criteria_col else ""

            parsed_date = parse_date(date_str) or date_str

            milestone = {
                "name": name,
                "date": parsed_date,
                "criteria": criteria,
                "category": classify_milestone(name),
            }
            milestones.append(milestone)

    return milestones


# ============================================================
# 关键路径解析
# ============================================================

def parse_critical_paths(paragraphs: list) -> list:
    """从段落中提取关键路径信息

    关键路径在文档中以 "关键路径X-名称" 开头，
    后续段落中包含设计/采购/施工各阶段的时间节点。
    使用正则提取每个带日期的关键步骤。
    """
    paths = []
    current_path = None
    collecting = False

    for p in paragraphs:
        text = p.strip()
        if not text:
            continue

        # 检测关键路径标题行（如 "关键路径1-原料粉磨车间..."）
        m = re.match(r"关键路径\s*(\d+)[：:\-—]?\s*(.+)", text)
        if m:
            if current_path:
                paths.append(current_path)
            current_path = {
                "name": f"关键路径{m.group(1)}-{m.group(2).strip()[:120]}",
                "steps": [],
            }
            collecting = True
            continue

        # 如果遇到下一个大标题（不含"关键路径"但像是新的章节标题），停止收集
        if collecting and current_path is not None and re.match(r"^\d+\.\d+", text):
            # 是编号标题（如 5.1.3），停止当前路径收集
            collecting = False
            continue

        # 收集当前关键路径下的内容
        if collecting and current_path is not None:
            # 提取所有日期模式的行：
            # 格式1: "设计：图纸12.29前完成"
            # 格式2: "2021.4.10开始..."
            # 格式3: "4.10前完成..."
            date_pattern = re.compile(
                r"(\d{4}[.\-/]\d{1,2}[.\-/]\d{1,2}(?:前)?"  # 完整日期 2021.4.10前
                r"|\d{1,2}[.\-/]\d{1,2}(?:前|前完成|前提供|前开始|日前|日)"  # M.D前 4.10前
                r"|\d{4}年\d{1,2}月\d{1,2}日"  # 中文日期
                r"|\d{1,2}月\d{1,2}日)"  # M月D日
            )

            # 按句子拆分
            sentences = re.split(r"[；;。]", text)
            for sent in sentences:
                sent = sent.strip()
                if not sent or len(sent) < 8:
                    continue

                dates = date_pattern.findall(sent)
                if dates:
                    # 取第一个日期
                    date_str = dates[0]
                    full_date_str = date_str
                    # 补全年份（如果日期不含年份，从上下文推断：文档在前50段提到2020-2021）
                    if not re.match(r"\d{4}", date_str):
                        # M.D格式，推断年份为2021
                        m2 = re.match(r"(\d{1,2})[.\-/](\d{1,2})", date_str)
                        if m2:
                            mo, d = int(m2.group(1)), int(m2.group(2))
                            # 如果月份是10-12月，可能是2020年；否则2021年
                            y = 2020 if mo >= 10 else 2021
                            full_date_str = f"{y}-{mo:02d}-{d:02d}"
                        else:
                            full_date_str = date_str

                    parsed_date = parse_date(full_date_str) or full_date_str
                    phase = classify_path_step(sent)

                    current_path["steps"].append({
                        "phase": phase,
                        "description": sent.strip()[:200],
                        "date": parsed_date,
                    })

    if current_path:
        paths.append(current_path)

    return paths


# ============================================================
# 月度任务解析
# ============================================================

def parse_monthly_tasks(tables_data: list) -> list:
    """从表格中解析月度任务分解

    月度计划表格结构（Table 2）：
    - Row0: 表头（6列：月度 | 里程碑事件 x5）
    - Row1: 子表头（5列：月度 | 设计 | 采购发货(采购部) | 采购发货(装备所) | 施工(土建)）
    - Row2+: 数据行（5列，第6列"施工安装"合并到第5列）
    """
    monthly = []

    for rows in tables_data:
        if not rows or len(rows) < 4:
            continue

        header = rows[0]

        # 识别月度任务表格：第1列表头含"月度"
        has_month = any("月度" in c for c in header)
        is_wide = len(header) >= 4 or (len(rows) > 1 and len(rows[1]) >= 4)

        if not (has_month and is_wide):
            continue

        # 使用位置映射（文档表格结构固定）
        # Col0=月度, Col1=设计, Col2=采购国内, Col3=采购国际/装备, Col4=施工土建, Col5=施工安装（可选）
        month_col = 0
        design_col = 1 if len(header) > 1 else None
        proc_dom_col = 2 if len(header) > 2 else None
        proc_intl_col = 3 if len(header) > 3 else None
        const_civil_col = 4 if len(header) > 4 else None
        const_install_col = 5 if len(header) > 5 else None

        # 从第2行开始解析（跳过Row0表头和Row1子表头）
        for row in rows[2:]:
            if len(row) < 5:  # 至少需要5列数据
                continue

            month_label = row[month_col].strip() if month_col < len(row) else ""
            if not month_label:
                continue

            # 跳过子表头重复行
            if "月度" in month_label or "设计" in month_label or "里程碑" in month_label:
                continue

            month_date = parse_date(month_label) or ""

            entry = {
                "month": month_date,
                "label": month_label[:50],
                "design": _extract_task_items(row[design_col]) if design_col is not None and design_col < len(row) else [],
                "procurement_domestic": _extract_task_items(row[proc_dom_col]) if proc_dom_col is not None and proc_dom_col < len(row) else [],
                "procurement_international": _extract_task_items(row[proc_intl_col]) if proc_intl_col is not None and proc_intl_col < len(row) else [],
                "construction_civil": _extract_task_items(row[const_civil_col]) if const_civil_col is not None and const_civil_col < len(row) else [],
                "construction_install": _extract_task_items(row[const_install_col]) if const_install_col is not None and const_install_col < len(row) else [],
            }
            monthly.append(entry)

    return monthly


def _extract_task_items(cell_text: str) -> list:
    """从单元格中提取单个任务项（按换行或 * 或数字编号分割）"""
    if not cell_text or not cell_text.strip():
        return []

    text = cell_text.strip()
    # 按 * 前缀或换行分割
    items = re.split(r"\n|\*\s*", text)
    result = []
    for item in items:
        item = item.strip()
        if item and len(item) > 2:
            result.append(item[:200])
    return result


# ============================================================
# 主解析入口
# ============================================================

# ============================================================
# .doc 格式转换
# ============================================================

def _find_libreoffice() -> Optional[str]:
    """查找 LibreOffice 可执行文件路径"""
    # Windows 常见路径
    candidates = [
        "soffice",
        "soffice.exe",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for path in candidates:
        if shutil.which(path) or os.path.isfile(path):
            return path
    # 尝试从注册表查找（Windows）
    try:
        import winreg
        for root in [winreg.HKEY_LOCAL_MACHINE, winreg.HKEY_CURRENT_USER]:
            for sub in [r"SOFTWARE\LibreOffice", r"SOFTWARE\WOW6432Node\LibreOffice"]:
                try:
                    with winreg.OpenKey(root, sub) as key:
                        return os.path.join(winreg.QueryValueEx(key, "Path")[0], "program", "soffice.exe")
                except OSError:
                    pass
    except Exception:
        pass
    return None


def _convert_doc_to_docx_bytes(file_bytes: bytes, filename: str) -> bytes:
    """将 .doc 文件转换为 .docx 格式

    尝试方案（按优先级）：
    1. LibreOffice headless 转换（最可靠，保留表格）
    2. python-docx 直接打开（有些 .doc 实际是 .docx 格式）
    3. 二进制文本提取（兜底方案）
    """
    # 方案1: 先尝试直接用 python-docx 打开（可能是伪装的 .docx）
    try:
        doc = Document(BytesIO(file_bytes))
        # 如果段落数/表格数 > 0，说明确实是 docx 格式
        if len(doc.paragraphs) > 0 or len(doc.tables) > 0:
            return file_bytes
    except Exception:
        pass

    # 方案2: 使用 LibreOffice 转换
    libreoffice = _find_libreoffice()
    if libreoffice:
        tmpdir = tempfile.mkdtemp()
        try:
            # 写入原始 .doc 文件
            doc_path = os.path.join(tmpdir, filename)
            with open(doc_path, "wb") as f:
                f.write(file_bytes)

            # 调用 LibreOffice 转换
            result = subprocess.run(
                [libreoffice, "--headless", "--convert-to", "docx", "--outdir", tmpdir, doc_path],
                capture_output=True, text=True, timeout=60,
                env={**os.environ, "HOME": tmpdir}
            )

            # 查找生成的 .docx 文件
            base = os.path.splitext(filename)[0]
            out_path = os.path.join(tmpdir, base + ".docx")
            if os.path.isfile(out_path):
                with open(out_path, "rb") as f:
                    return f.read()

            # 可能有其他命名
            for fname in os.listdir(tmpdir):
                if fname.endswith(".docx"):
                    with open(os.path.join(tmpdir, fname), "rb") as f:
                        return f.read()
        except (subprocess.TimeoutExpired, FileNotFoundError) as e:
            pass
        finally:
            try:
                shutil.rmtree(tmpdir, ignore_errors=True)
            except Exception:
                pass

    # 方案3: 尝试 antiword 命令行工具
    try:
        result = subprocess.run(
            ["antiword", "-"], input=file_bytes, capture_output=True, timeout=30
        )
        text = result.stdout.decode("utf-8", errors="ignore")
        if text and len(text.strip()) > 100:
            # 将纯文本包装成最小化 .docx
            from docx import Document as NewDoc
            doc = NewDoc()
            for line in text.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
            buf = BytesIO()
            doc.save(buf)
            return buf.getvalue()
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # 方案4: 无法转换，给出明确错误
    raise RuntimeError(
        "无法解析 .doc 格式文件。\n"
        "请使用以下任一方式：\n"
        "1. 用 Microsoft Word 或 WPS 打开文件，另存为 .docx 格式后上传\n"
        "2. 安装 LibreOffice 后重试（支持自动转换）\n"
        "3. 直接上传已有的 .docx 版本文件"
    )


def _extract_text_from_binary_doc(file_bytes: bytes) -> bytes:
    """从 OLE2 (.doc) 二进制文件中尽可能提取文本

    将提取的文本包装成简化版 .docx XML 格式，
    以便现有解析流程能处理。
    """
    text = _extract_raw_text_from_ole(file_bytes)

    if not text or len(text.strip()) < 50:
        raise RuntimeError(
            "无法读取 .doc 文件内容。请使用 Microsoft Word 或 WPS 将文件另存为 .docx 格式后重新上传。"
        )

    # 将纯文本包装成最小化的 .docx
    # python-docx 可以读取这个简单结构
    from docx import Document as NewDoc
    doc = NewDoc()
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _extract_raw_text_from_ole(file_bytes: bytes) -> str:
    """从 OLE2 容器中提取可读文本"""
    try:
        import olefile
    except ImportError:
        # 无 olefile，用简单二进制扫描
        return _scan_binary_for_text(file_bytes)

    try:
        ole = olefile.OleFileIO(file_bytes)
        # Word文档文本通常在 WordDocument 流中
        # 尝试读取主文本流
        text_parts = []

        # 列出所有流
        for stream_name in ole.listdir():
            stream_path = "/".join(stream_name)
            try:
                data = ole.openstream(stream_name).read()
                # 尝试解码为 UTF-16LE
                try:
                    decoded = data.decode("utf-16-le", errors="ignore")
                    # 过滤出可读的中英文文本片段
                    readable = "".join(
                        c for c in decoded
                        if c.isprintable() or c in "\n\r\t"
                    )
                    if len(readable) > 100:
                        text_parts.append(readable)
                except Exception:
                    pass
            except Exception:
                pass

        ole.close()
        return "\n".join(text_parts)
    except Exception:
        return _scan_binary_for_text(file_bytes)


def _scan_binary_for_text(file_bytes: bytes) -> str:
    """从二进制数据中扫描提取可读的中英文文本"""
    # 尝试 UTF-16LE 解码
    try:
        text = file_bytes.decode("utf-16-le", errors="ignore")
    except Exception:
        text = file_bytes.decode("latin-1", errors="ignore")

    # 过滤出有效的文本段落（至少包含2个中文字符或足够长度的英文单词）
    lines = []
    current = []
    for ch in text:
        if ch.isprintable() or ch in "\t":
            current.append(ch)
        else:
            if current:
                line = "".join(current).strip()
                if len(line) > 3 and _has_readable_content(line):
                    lines.append(line)
                current = []
    if current:
        line = "".join(current).strip()
        if len(line) > 3:
            lines.append(line)

    return "\n".join(lines)


def _has_readable_content(text: str) -> bool:
    """检查文本是否包含有意义的内容（中文或英文单词）"""
    chinese_chars = len(re.findall(r"[一-鿿]", text))
    alpha_chars = len(re.findall(r"[a-zA-Z]", text))
    return chinese_chars >= 2 or alpha_chars >= 5


# ============================================================
# 主解析入口
# ============================================================

def parse_planning_document(file_bytes: bytes, filename: str) -> dict:
    """解析策划书文档，返回完整结构化数据

    Args:
        file_bytes: 文档字节流
        filename: 原始文件名

    Returns:
        {
            "project": {...},
            "milestones": [...],
            "critical_paths": [...],
            "monthly_tasks": [...]
        }
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext not in (".docx", ".doc"):
        raise RuntimeError(f"不支持的文件格式: {ext}，请上传 .docx 或 .doc 文件")

    # .doc 格式先转换为 .docx
    if ext == ".doc":
        try:
            file_bytes = _convert_doc_to_docx_bytes(file_bytes, filename)
        except RuntimeError:
            raise
        except Exception as e:
            raise RuntimeError(f".doc 文件转换失败: {str(e)}")

    try:
        doc = Document(BytesIO(file_bytes))
    except Exception as e:
        raise RuntimeError(f"无法打开文档: {str(e)}。请确保上传的是有效的 .docx 或 .doc 文件。")

    # 提取段落文本
    paragraphs = [p.text for p in doc.paragraphs]

    # 提取表格
    tables_data = extract_tables(doc)

    # 解析各部分
    project = parse_project_info(paragraphs)
    milestones = parse_milestones(tables_data)
    critical_paths = parse_critical_paths(paragraphs)
    monthly_tasks = parse_monthly_tasks(tables_data)

    # 自动补全项目名称
    if not project.get("name"):
        # 从文档前几段拼接
        name_lines = [p.strip() for p in paragraphs[:10] if p.strip() and len(p.strip()) < 100]
        project["name"] = " ".join(name_lines[:3]) if name_lines else filename

    return {
        "project": project,
        "milestones": milestones,
        "critical_paths": critical_paths,
        "monthly_tasks": monthly_tasks,
    }
